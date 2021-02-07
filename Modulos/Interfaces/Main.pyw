from tkinter import *
from tkinter.ttk import *
from docx import *
from docx.shared import Pt


class Control():
    def __init__(self, tk_root):
        tk_root.geometry("600x620")
        tk_root.config(bg="#292929")
        tk_root.title("Dialogs-M")
        self.menuBar(tk_root)

        MainFrame(tk_root,
            self.mainFrameS(),
            self.MLabelFrameS(),
            self.shortFrameS()
            )

    def mainFrameS(self):
        MFS = Style()
        MFS.configure('MFS.TFrame',
            background="#ffffff"
            )
        return MFS

    def shortFrameS(self):
        SFS = Style()
        SFS.configure('SFS.TFrame',
            background="#363c3f"
            )
        return SFS

    def MLabelFrameS(self):
        LFS = Style()
        LFS.configure('lfs.TLabelframe',
        background="#ffffff"
            )

        LFS.configure('lfs.TLabelframe.Label',
            Foreground="#ffffff",
            background="#ffffff",
            font=("cambria", 13)
            )

        return LFS

    def menuBar(self, parent):
        self.menuBar = Menu(parent)
        parent.config(menu=self.menuBar)
        self.archivoMenu = Menu(self.menuBar, tearoff=0)
        self.menuBar.add_cascade(label="Archivo", menu=self.archivoMenu)
        self.archivoMenu.add_command(label="Nuevo", command=lambda:self.Nuevo())
        self.archivoMenu.add_command(label="Guardar", command=lambda:self.Guardar())
        self.archivoMenu.add_command(label="Guardar como", command=lambda:self.SaveAs())
        self.archivoMenu.add_separator()
        self.archivoMenu.add_command(label="Salir", command=lambda:self.Salir())
        self.preMenu = Menu(self.menuBar, tearoff=0)
        self.menuBar.add_cascade(label= "Preferencias", menu=self.preMenu)
        self.preMenu.add_command(label="Temas", command=lambda:self.temas())
        self.preMenu.add_command(label="Fondos", command=lambda:self.imagen())


class MainFrame():
    def __init__(self, root, Fstyle, Lstyle, Sstyle):
        self.root = root
        self.main()
        self.labelFrame()
        self.shortFrame()
        Inside(self.root,)

    def shortFrame(self):
        short_Frame = Frame(self.MFrame, width=600, height=69, style='SFS.TFrame')
        short_Frame.place(x=0, y=0)

    def main(self):
        self.MFrame = Frame(
            self.root,
            width=580,
            height=600,
            style="MFS.TFrame"
            )
        self.MFrame.place(x=10, y=10)

    def labelFrame(self):
            self.MLabelFrame = LabelFrame(
                self.MFrame,
                text="Datos del dialogo",
                width=560, heigh=150,
                style="lfs.TLabelframe"
                )

            self.MLabelFrame.place(y=150, x=10)

            self.DialogLabelFrame = LabelFrame(
                self.MFrame,
                text="Dialogo",
                width=560, heigh=230,
                style="lfs.TLabelframe"
                )

            self.DialogLabelFrame.place(y=320, x=10)


class Inside():

    def __init__(self, Mframe):
        self.mainFrameL(Mframe)
        self.mainFrameT(Mframe)
        GettingText(Mframe, self.DocName, self.DialogPlace, self.PjsName, self.PjsState, self.DialogEmotion, self.Dialog)

    def mainFrameL(self, parent):

        Label(
            parent,
            text="Rellena los siguientes espacios con la informacion correspondiente.",
            font=("segoe ui", 13),
            background="#363c3f",
            foreground="#ffffff"
            ).place(x=20, y=25)

        Label(
            parent,
            text="Ingresa el nombre del documento aqui.",
            font=("segoe ui", 14),
            background="#ffffff",
            foreground="#000000"
            ).place(x=20, y=80)


        Label(
            parent,
            text="Lugar :",
            font=("segoe ui", 14),
            background="#ffffff",
            foreground="#000000"
            ).place(x=30, y=190)

        Label(
            parent,
            text="Nombre :",
            font=("segoe ui", 14),
            background="#ffffff",
            foreground="#000000"
            ).place(x=320, y=190)

        Label(
            parent,
            text="Emocion :",
            font=("segoe ui", 14),
            background="#ffffff",
            foreground="#000000"
            ).place(x=30, y=250)

        Label(
            parent,
            text="Tipo :",
            font=("segoe ui", 14),
            background="#ffffff",
            foreground="#000000"
            ).place(x=320, y=250)

    def mainFrameT(self, parent):

        self.DocName = Entry(
            parent,
            font=("segoe ui", 14),
            width=40,
            )
        self.DocName.place(x=20, y=120)

        self.DialogPlace = Entry(
            parent,
            font=("segoe ui", 12),
            width=22
            )
        self.DialogPlace.place(x=97, y=192)

        self.PjsName = Entry(
            parent,
            font=("segoe ui", 12),
            width=17
            )
        self.PjsName.place(x=410, y=192)

        self.PjsState = Entry(
            parent,
            font=("segoe ui", 12),
            width=17
            )
        self.PjsState.place(x=120, y=252)

        self.DialogEmotion = Entry(
            parent,
            font=("segoe ui", 12),
            width=17
            )
        self.DialogEmotion.place(x=378, y=252)

        self.Dialog = Text(
            parent,
            font=("segoe ui", 12),
            width=60,
            height=9
            )
        self.Dialog.place(x=30, y=360)


class GettingText():
    def __init__(self, parent, title, place, name, emotion, type, dialog):

        self.buttons(parent)
        self.title = title
        self.place = place
        self.name = name
        self.emotion = emotion
        self.type = type
        self.dialog = dialog

    def generateDoc(self):
        document = Document()
        document.add_heading(self.title.get(), 0)
        #fontStuf
        run = document.add_paragraph().add_run()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Cambria'
        font.size = Pt(15)
        #paragraph
        document.add_paragraph("Tipo de dialogo: " + self.type.get()+"\n")
        document.add_paragraph("Lugar: " + self.place.get()+"\n")
        document.add_paragraph("Emocion del emisor: " + self.emotion.get()+"\n")
        document.add_paragraph("Dialogo de " + "\\n["+self.name.get()+"]"+ ": " + self.dialog.get('1.0', 'end'))
        document.save('..\\'+self.title.get()+'.docx')

    def buttons(self, parent):

        """self.newDialog = Button(
            parent,
            text="Nuevo dialogo",
            width=25,
            command=lambda:self.titulo()
            ).place(x=20, y=570)"""

        self.save = Button(
            parent,
            text="Guardar",
            width=15,
            command=lambda:self.generateDoc(),
            ).place(x=450, y=570)

        """self.saveOther = Button(
            parent,
            text="Guardar y otro",
            width=20
            ).place(x=450, y=570)"""
